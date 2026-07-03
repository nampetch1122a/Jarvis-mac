import os
import re
import json
import base64
import anthropic
import fitz
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from datetime import datetime

app = Flask(__name__)
CORS(app)
client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# --- Paths: อยู่ในโปรเจกต์เอง ไม่ใช้ home directory ของเครื่อง ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
OUTPUT_DIR = os.path.join(DATA_DIR, "outputs")
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
KNOWLEDGE_DIR = os.path.join(DATA_DIR, "knowledge")
STATE_FILE = os.path.join(DATA_DIR, "jarvis_state.json")

for d in (OUTPUT_DIR, UPLOAD_DIR, KNOWLEDGE_DIR):
    os.makedirs(d, exist_ok=True)

if not os.path.exists(STATE_FILE):
    with open(STATE_FILE, "w") as f:
        json.dump({"status": "idle", "text": ""}, f, ensure_ascii=False)


def extract_json(raw):
    """ดึง JSON object ออกจากคำตอบ AI แม้จะมีข้อความอธิบาย/ขั้นตอนคิดปนมาด้วยก็ตาม"""
    text = re.sub(r'```json|```', '', raw).strip()
    start = text.find('{')
    if start == -1:
        raise ValueError("ไม่พบ JSON ในคำตอบ")
    depth = 0
    for i in range(start, len(text)):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return json.loads(text[start:i+1])
    raise ValueError("JSON ไม่สมบูรณ์ (วงเล็บปีกกาไม่ปิด)")


def load_knowledge():
    knowledge = ""
    try:
        for fname in os.listdir(KNOWLEDGE_DIR):
            fpath = os.path.join(KNOWLEDGE_DIR, fname)
            if fname.endswith(".txt") or fname.endswith(".csv"):
                with open(fpath, encoding="utf-8-sig") as f:
                    knowledge += f"\n\n=== {fname} ===\n" + f.read()
            elif fname.endswith(".pdf"):
                doc = fitz.open(fpath)
                text = "\n".join([page.get_text() for page in doc])
                knowledge += f"\n\n=== {fname} ===\n" + text
    except Exception as e:
        print(f"Knowledge load error: {e}")
    return knowledge[:8000]


def calculate_from_csv(content):
    income_keywords = ["ขาย", "รายได้", "รับชำระ", "บริการ"]
    expense_keywords = ["เงินเดือน", "ค่าเช่า", "ค่าไฟ", "ค่าขนส่ง", "ค่าการตลาด", "ค่าซ่อม", "ดอกเบี้ยจ่าย", "ค่าสาธารณูปโภค", "ค่าประกัน", "ค่าวัสดุ", "ค่าโฆษณา", "ค่าน้ำ", "ค่าแก๊ส", "ค่าบำรุง"]
    cogs_keywords = ["ซื้อวัตถุดิบ", "ซื้อสินค้า", "ต้นทุน", "วัตถุดิบ"]
    return_in_keywords = ["รับคืน", "สินค้ารับคืน"]
    return_out_keywords = ["ส่งคืน", "สินค้าส่งคืน", "คืนสินค้า"]
    skip_keywords = ["ทุนจดทะเบียน", "ทุน", "กู้ยืม", "เงินกู้", "ที่ดิน"]

    total_income = 0.0
    total_expense = 0.0
    total_cogs = 0.0
    total_return_in = 0.0
    total_return_out = 0.0

    lines = content.split("\n")
    for line in lines[1:]:
        cols = [c.strip() for c in line.split(",")]
        if len(cols) < 3:
            continue
        desc = cols[1] if len(cols) > 1 else ""
        try:
            if len(cols) > 4 and cols[4].strip():
                amount = float(cols[4].replace(",", ""))
            else:
                debit = float(cols[2].replace(",", "")) if cols[2].strip() else 0
                credit = float(cols[3].replace(",", "")) if len(cols) > 3 and cols[3].strip() else 0
                amount = debit or credit
        except Exception:
            continue
        if amount == 0:
            continue
        if any(k in desc for k in skip_keywords):
            continue
        if any(k in desc for k in return_in_keywords):
            total_return_in += amount
        elif any(k in desc for k in return_out_keywords):
            total_return_out += amount
        elif any(k in desc for k in income_keywords):
            total_income += amount
        elif any(k in desc for k in cogs_keywords):
            total_cogs += amount
        elif any(k in desc for k in expense_keywords):
            total_expense += amount

    net_income = total_income - total_return_in
    net_expense = total_expense + total_cogs - total_return_out
    net_profit = net_income - net_expense
    profit_margin = (net_profit / net_income * 100) if net_income > 0 else 0

    return {
        "total_income": net_income,
        "total_expense": net_expense,
        "net_profit": net_profit,
        "profit_margin": round(profit_margin, 1),
        "raw_income": total_income,
        "raw_expense": total_expense + total_cogs,
        "return_in": total_return_in,
        "return_out": total_return_out
    }


def read_pdf(path):
    doc = fitz.open(path)
    return "\n".join([page.get_text() for page in doc])


def read_docx(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])


def read_excel(path):
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append("\t".join([str(c) if c is not None else "" for c in row]))
    return "\n".join(rows)


def read_csv(path):
    with open(path, encoding="utf-8-sig") as f:
        return f.read()


def clean_for_speech(text):
    text = re.sub(r'[#*_`|•→\[\]()]', '', text)
    text = re.sub(r'-{2,}', '', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def title_fill():
    return PatternFill("solid", fgColor="1F4E79")


def section_fill():
    return PatternFill("solid", fgColor="BDD7EE")


def subtotal_fill():
    return PatternFill("solid", fgColor="DEEAF1")


def thin_border():
    s = Side(style="thin", color="AAAAAA")
    return Border(left=s, right=s, top=s, bottom=s)


def set_col_width(ws, col, width):
    ws.column_dimensions[get_column_letter(col)].width = width


def write_cell(ws, row, col, value, bold=False, fill=None, align="left", num_fmt=None, font_color="000000", size=11):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = Font(bold=bold, color=font_color, size=size)
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    cell.border = thin_border()
    if fill:
        cell.fill = fill
    if num_fmt:
        cell.number_format = num_fmt
    return cell


def create_excel_report(filename, data, ai_analysis, company_name="บริษัท"):
    wb = openpyxl.Workbook()

    ws1 = wb.active
    ws1.title = "งบกำไรขาดทุน"
    set_col_width(ws1, 1, 45)
    set_col_width(ws1, 2, 20)
    r = 1
    ws1.merge_cells(f"A{r}:B{r}")
    write_cell(ws1, r, 1, company_name, bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=14)
    r += 1
    ws1.merge_cells(f"A{r}:B{r}")
    write_cell(ws1, r, 1, "งบกำไรขาดทุนเบ็ดเสร็จ (Income Statement)", bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=12)
    r += 1
    ws1.merge_cells(f"A{r}:B{r}")
    write_cell(ws1, r, 1, f"สำหรับงวด สิ้นสุดวันที่ {datetime.now().strftime('%d/%m/%Y')}", fill=subtotal_fill(), align="center")
    r += 1
    ws1.merge_cells(f"A{r}:B{r}")
    write_cell(ws1, r, 1, "หน่วย: บาท", fill=subtotal_fill(), align="center")
    r += 1
    write_cell(ws1, r, 1, "รายการ", bold=True, fill=section_fill(), align="center")
    write_cell(ws1, r, 2, "จำนวนเงิน (บาท)", bold=True, fill=section_fill(), align="center")
    r += 1
    rows_income = [
        ("รายได้จากการขาย", data.get("raw_income", 0), False),
        ("  หัก สินค้ารับคืน", -data.get("return_in", 0), False),
        ("รายได้สุทธิ", data.get("total_income", 0), True),
        ("", None, False),
        ("  หัก ต้นทุนสินค้า", -data.get("raw_expense", 0), False),
        ("  บวก สินค้าส่งคืน", data.get("return_out", 0), False),
        ("ต้นทุนสุทธิ", -data.get("total_expense", 0), True),
        ("", None, False),
        ("กำไรขั้นต้น (Gross Profit)", data.get("total_income", 0) - data.get("total_expense", 0), True),
        ("", None, False),
        ("กำไร (ขาดทุน) สุทธิ (Net Profit)", data.get("net_profit", 0), True),
        ("อัตรากำไรสุทธิ (%)", data.get("profit_margin", 0), False),
    ]
    for label, val, bold in rows_income:
        if val is None:
            ws1.cell(row=r, column=1).value = ""
            ws1.cell(row=r, column=2).value = ""
            r += 1
            continue
        f = subtotal_fill() if bold else None
        write_cell(ws1, r, 1, label, bold=bold, fill=f)
        if "%" in label:
            write_cell(ws1, r, 2, val / 100, bold=bold, fill=f, align="right", num_fmt='0.00%')
        else:
            write_cell(ws1, r, 2, val, bold=bold, fill=f, align="right", num_fmt='#,##0.00')
        r += 1

    ws2 = wb.create_sheet("งบแสดงฐานะการเงิน")
    set_col_width(ws2, 1, 45)
    set_col_width(ws2, 2, 20)
    r = 1
    ws2.merge_cells(f"A{r}:B{r}")
    write_cell(ws2, r, 1, company_name, bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=14)
    r += 1
    ws2.merge_cells(f"A{r}:B{r}")
    write_cell(ws2, r, 1, "งบแสดงฐานะการเงิน (Balance Sheet)", bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=12)
    r += 1
    ws2.merge_cells(f"A{r}:B{r}")
    write_cell(ws2, r, 1, f"ณ วันที่ {datetime.now().strftime('%d/%m/%Y')}", fill=subtotal_fill(), align="center")
    r += 1
    ws2.merge_cells(f"A{r}:B{r}")
    write_cell(ws2, r, 1, "หน่วย: บาท", fill=subtotal_fill(), align="center")
    r += 1
    write_cell(ws2, r, 1, "รายการ", bold=True, fill=section_fill(), align="center")
    write_cell(ws2, r, 2, "จำนวนเงิน (บาท)", bold=True, fill=section_fill(), align="center")
    r += 1
    rows_bs = [
        ("สินทรัพย์ (Assets)", None, True, title_fill(), "FFFFFF"),
        ("สินทรัพย์หมุนเวียน (Current Assets)", None, True, section_fill(), "000000"),
        ("  เงินสดและรายการเทียบเท่า", data.get("cash_and_equivalents", 0), False, None, "000000"),
        ("  ลูกหนี้การค้า", data.get("accounts_receivable", 0), False, None, "000000"),
        ("  สินค้าคงเหลือ", data.get("inventory", 0), False, None, "000000"),
        ("  ค่าใช้จ่ายจ่ายล่วงหน้า", data.get("prepaid_expenses", 0), False, None, "000000"),
        ("  รวมสินทรัพย์หมุนเวียน", data.get("current_assets", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("สินทรัพย์ไม่หมุนเวียน (Non-Current Assets)", None, True, section_fill(), "000000"),
        ("  ที่ดิน อาคาร และอุปกรณ์", data.get("fixed_assets", 0), False, None, "000000"),
        ("  รวมสินทรัพย์ไม่หมุนเวียน", data.get("fixed_assets", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("รวมสินทรัพย์ทั้งหมด", data.get("total_assets", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("หนี้สินและส่วนของผู้ถือหุ้น", None, True, title_fill(), "FFFFFF"),
        ("หนี้สินหมุนเวียน (Current Liabilities)", None, True, section_fill(), "000000"),
        ("  เจ้าหนี้การค้า", data.get("accounts_payable", 0), False, None, "000000"),
        ("  ค่าใช้จ่ายค้างจ่าย", data.get("accrued_expenses", 0), False, None, "000000"),
        ("  รวมหนี้สินหมุนเวียน", data.get("current_liabilities", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("หนี้สินไม่หมุนเวียน (Non-Current Liabilities)", None, True, section_fill(), "000000"),
        ("  เงินกู้ระยะยาว", data.get("long_term_debt", 0), False, None, "000000"),
        ("  รวมหนี้สินไม่หมุนเวียน", data.get("long_term_debt", 0), True, subtotal_fill(), "000000"),
        ("  รวมหนี้สินทั้งหมด", data.get("total_liabilities", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("ส่วนของผู้ถือหุ้น (Equity)", None, True, section_fill(), "000000"),
        ("  ทุนจดทะเบียน", data.get("paid_in_capital", 0), False, None, "000000"),
        ("  กำไรสะสม", data.get("retained_earnings", 0), False, None, "000000"),
        ("  รวมส่วนของผู้ถือหุ้น", data.get("equity", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("รวมหนี้สินและส่วนของผู้ถือหุ้น", data.get("total_liabilities_equity", 0), True, subtotal_fill(), "000000"),
    ]
    for label, val, bold, fill, fc in rows_bs:
        if val is None and not bold:
            ws2.cell(row=r, column=1).value = ""
            ws2.cell(row=r, column=2).value = ""
            r += 1
            continue
        write_cell(ws2, r, 1, label, bold=bold, fill=fill, font_color=fc)
        if val is not None:
            write_cell(ws2, r, 2, val, bold=bold, fill=fill, align="right", num_fmt='#,##0.00', font_color=fc)
        else:
            write_cell(ws2, r, 2, "", fill=fill, font_color=fc)
        r += 1

    ws3 = wb.create_sheet("งบกระแสเงินสด")
    set_col_width(ws3, 1, 45)
    set_col_width(ws3, 2, 20)
    r = 1
    ws3.merge_cells(f"A{r}:B{r}")
    write_cell(ws3, r, 1, company_name, bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=14)
    r += 1
    ws3.merge_cells(f"A{r}:B{r}")
    write_cell(ws3, r, 1, "งบกระแสเงินสด (Cash Flow Statement)", bold=True, fill=title_fill(), align="center", font_color="FFFFFF", size=12)
    r += 1
    ws3.merge_cells(f"A{r}:B{r}")
    write_cell(ws3, r, 1, f"สำหรับงวด สิ้นสุดวันที่ {datetime.now().strftime('%d/%m/%Y')}", fill=subtotal_fill(), align="center")
    r += 1
    ws3.merge_cells(f"A{r}:B{r}")
    write_cell(ws3, r, 1, "หน่วย: บาท", fill=subtotal_fill(), align="center")
    r += 1
    write_cell(ws3, r, 1, "รายการ", bold=True, fill=section_fill(), align="center")
    write_cell(ws3, r, 2, "จำนวนเงิน (บาท)", bold=True, fill=section_fill(), align="center")
    r += 1
    rows_cf = [
        ("กิจกรรมดำเนินงาน (Operating Activities)", None, True, section_fill(), "000000"),
        ("  เงินสดรับจากลูกค้า", data.get("operating_cash_in", 0), False, None, "000000"),
        ("  เงินสดจ่ายค่าใช้จ่ายดำเนินงาน", -data.get("operating_cash_out", 0), False, None, "000000"),
        ("  กระแสเงินสดสุทธิจากดำเนินงาน", data.get("operating_cash_flow", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("กิจกรรมลงทุน (Investing Activities)", None, True, section_fill(), "000000"),
        ("  ซื้อสินทรัพย์ถาวร", -data.get("fixed_assets", 0), False, None, "000000"),
        ("  กระแสเงินสดสุทธิจากลงทุน", -data.get("fixed_assets", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("กิจกรรมจัดหาเงิน (Financing Activities)", None, True, section_fill(), "000000"),
        ("  เงินสดรับจากเงินกู้", data.get("long_term_debt", 0), False, None, "000000"),
        ("  เงินสดรับจากทุน", data.get("paid_in_capital", 0), False, None, "000000"),
        ("  กระแสเงินสดสุทธิจากจัดหาเงิน", data.get("financing_cash_flow", 0), True, subtotal_fill(), "000000"),
        ("", None, False, None, "000000"),
        ("กระแสเงินสดสุทธิทั้งหมด", data.get("net_cash", 0), True, subtotal_fill(), "000000"),
        ("เงินสดต้นงวด", 0, False, None, "000000"),
        ("เงินสดปลายงวด", data.get("cash_and_equivalents", 0), True, subtotal_fill(), "000000"),
    ]
    for label, val, bold, fill, fc in rows_cf:
        if val is None and not bold:
            ws3.cell(row=r, column=1).value = ""
            ws3.cell(row=r, column=2).value = ""
            r += 1
            continue
        write_cell(ws3, r, 1, label, bold=bold, fill=fill, font_color=fc)
        if val is not None:
            write_cell(ws3, r, 2, val, bold=bold, fill=fill, align="right", num_fmt='#,##0.00', font_color=fc)
        else:
            write_cell(ws3, r, 2, "", fill=fill, font_color=fc)
        r += 1

    ws4 = wb.create_sheet("บทวิเคราะห์ AI")
    set_col_width(ws4, 1, 80)
    r = 1
    write_cell(ws4, r, 1, "บทวิเคราะห์จาก Jarvis AI", bold=True, fill=title_fill(), font_color="FFFFFF", size=13)
    r += 2
    for line in ai_analysis.split("\n"):
        if line.strip():
            ws4.cell(row=r, column=1).value = line.strip()
            ws4.cell(row=r, column=1).alignment = Alignment(wrap_text=True)
            r += 1

    out_name = f"งบการเงิน_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    wb.save(out_path)
    return out_name


def create_exercise_report(filename, full_answer):
    doc = Document()
    title = doc.add_heading('เฉลยโจทย์บัญชี — โดย Jarvis AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"ไฟล์ต้นฉบับ: {filename}")
    doc.add_paragraph(f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    doc.add_heading('เฉลยและวิธีคำนวณ', level=1)
    for line in full_answer.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    out_name = f"เฉลย_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_name


def create_balance_sheet_report(filename, full_answer):
    doc = Document()
    title = doc.add_heading('งบแสดงฐานะการเงิน — โดย Jarvis AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"ไฟล์ต้นฉบับ: {filename}")
    doc.add_paragraph(f"วันที่: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    for line in full_answer.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    out_name = f"งบฐานะ_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    return out_name


def write_state(text, filename=""):
    try:
        with open(STATE_FILE, "w") as f:
            json.dump({"status": "responded", "text": text, "file": filename}, f, ensure_ascii=False)
    except Exception as e:
        print(f"State write error: {e}")


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/jarvis_state.json")
def jarvis_state():
    try:
        with open(STATE_FILE, encoding="utf-8") as f:
            return jsonify(json.load(f))
    except Exception:
        return jsonify({"status": "idle", "text": ""})


@app.route("/health")
def health():
    return jsonify({"status": "ok"})


@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    command = request.form.get("command", "สรุปเนื้อหาทั้งหมด")
    if not file:
        return jsonify({"error": "no file"}), 400
    path = os.path.join(UPLOAD_DIR, file.filename)
    file.save(path)
    ext = file.filename.lower().split(".")[-1]
    try:
        if ext == "pdf":
            content = read_pdf(path)
        elif ext == "docx":
            content = read_docx(path)
        elif ext in ("xlsx", "xls"):
            content = read_excel(path)
        elif ext == "csv":
            content = read_csv(path)
        else:
            return jsonify({"error": "ไฟล์ไม่รองรับ"}), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    knowledge = load_knowledge()
    want_excel = any(w in command for w in ["excel", "Excel", "xlsx", "สเปรดชีต", "ตาราง"])
    want_financial = any(w in command for w in ["งบการเงิน", "financial", "กำไรขาดทุน", "กระแสเงินสด", "ออกงบ"])
    want_solve = any(w in command for w in ["หาค่า", "เฉลย", "ตอบ", "แก้โจทย์", "หาตัวเลข", "คำนวณ", "หาค่าที่หายไป"])
    want_balance_sheet = any(w in command for w in ["งบฐานะ", "งบแสดงฐานะ", "balance sheet", "งบดุล", "ฐานะการเงิน"])
    want_word = any(w in command for w in ["word", "Word", "สรุปเป็น", "ออกเป็น", "ไฟล์", "รายงาน"]) or want_financial or want_solve or want_balance_sheet

    ai_data = {}
    docx_file = None
    answer = ""

    try:
        if want_excel and (want_financial or want_balance_sheet or "งบ" in command):
            prompt = f"""คุณคือ Jarvis AI ผู้เชี่ยวชาญด้านบัญชีมาตรฐานสากล

ความรู้จากแบบฝึกหัด:
{knowledge}

ข้อมูลจากไฟล์:
{content[:6000]}

กฎสำคัญ:
- รายได้ = ขาย + บริการ เท่านั้น ไม่รวมทุน/เงินกู้
- ค่าใช้จ่าย = เงินเดือน ค่าเช่า สาธารณูปโภค การตลาด ขนส่ง ซ่อมบำรุง ดอกเบี้ย
- สินทรัพย์ไม่หมุนเวียน = อุปกรณ์ อาคาร ที่ดิน
- หนี้สินไม่หมุนเวียน = เงินกู้ระยะยาว
- ส่วนของผู้ถือหุ้น = ทุนจดทะเบียน + กำไรสะสม
- total_assets ต้องเท่ากับ total_liabilities_equity

คำนวณและตรวจทานในใจให้เสร็จก่อน แล้วตอบกลับมาเป็น JSON object เพียงอย่างเดียวเท่านั้น ห้ามมีข้อความ คำอธิบาย หรือขั้นตอนการคิดใดๆ ปนอยู่ก่อนหรือหลัง JSON เด็ดขาด (ฟิลด์ "analysis" ต้องเขียนเป็นภาษาไทยเสมอ):
{{
  "company_name": "<ชื่อบริษัท>",
  "raw_income": <รายได้ก่อนหักคืน>,
  "return_in": <สินค้ารับคืน>,
  "total_income": <รายได้สุทธิ>,
  "raw_expense": <ค่าใช้จ่ายก่อนหักคืน>,
  "return_out": <สินค้าส่งคืน>,
  "total_expense": <ค่าใช้จ่ายสุทธิ>,
  "net_profit": <กำไรสุทธิ>,
  "profit_margin": <อัตรากำไร%>,
  "cash_and_equivalents": <เงินสดคงเหลือ>,
  "accounts_receivable": <ลูกหนี้การค้า>,
  "inventory": <สินค้าคงเหลือ>,
  "prepaid_expenses": <ค่าใช้จ่ายจ่ายล่วงหน้า>,
  "current_assets": <สินทรัพย์หมุนเวียนรวม>,
  "fixed_assets": <อุปกรณ์+อาคาร>,
  "total_assets": <สินทรัพย์รวม>,
  "accounts_payable": <เจ้าหนี้การค้า>,
  "accrued_expenses": <ค่าใช้จ่ายค้างจ่าย>,
  "current_liabilities": <หนี้สินหมุนเวียน>,
  "long_term_debt": <เงินกู้ระยะยาว>,
  "total_liabilities": <หนี้สินรวม>,
  "paid_in_capital": <ทุนจดทะเบียน>,
  "retained_earnings": <กำไรสะสม>,
  "equity": <ส่วนของผู้ถือหุ้น>,
  "total_liabilities_equity": <ต้องเท่ากับ total_assets>,
  "operating_cash_in": <เงินสดรับจากขาย/บริการ>,
  "operating_cash_out": <เงินสดจ่ายดำเนินงาน>,
  "operating_cash_flow": <กระแสเงินสดดำเนินงาน>,
  "financing_cash_flow": <เงินกู้+ทุน>,
  "net_cash": <กระแสเงินสดสุทธิ>,
  "analysis": "<วิเคราะห์ 5-6 ประโยค>"
}}"""

            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = message.content[0].text
            try:
                ai_data = extract_json(raw)
                company = ai_data.get("company_name", "บริษัทตัวอย่าง")
                analysis = ai_data.get("analysis", "")
                answer = clean_for_speech(analysis)
                docx_file = create_excel_report(file.filename, ai_data, analysis, company)
            except Exception as e:
                print(f"JSON parse error: {e}")
                answer = "ขออภัยครับ Jarvis ประมวลผลข้อมูลไม่สำเร็จ ลองส่งคำสั่งใหม่อีกครั้งครับ"
                docx_file = None

        elif want_balance_sheet:
            prompt = f"""คุณคือผู้เชี่ยวชาญบัญชีขั้นต้น
{knowledge}

ข้อมูล:
{content[:6000]}

จัดทำงบแสดงฐานะการเงินแบบรายงาน แยกหมวดครบ พร้อมหาค่าที่หายไปและแสดงวิธีคำนวณ ตอบภาษาไทย"""

            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2048,
                messages=[{"role": "user", "content": prompt}]
            )
            full_answer = message.content[0].text
            answer = clean_for_speech(full_answer[:500])
            docx_file = create_balance_sheet_report(file.filename, full_answer) if want_word else None

        elif want_solve:
            prompt = f"""คุณคือผู้เชี่ยวชาญบัญชีขั้นต้น
{knowledge}

โจทย์:
{content[:6000]}

กฎ:
1. ขายสุทธิ = ขาย - รับคืนและส่วนลด - ส่วนลดจ่าย
2. ต้นทุนขาย = สินค้าต้นงวด + ต้นทุนการซื้อสุทธิ - สินค้าปลายงวด
3. กำไรขั้นต้น = ขายสุทธิ - ต้นทุนขาย
4. ห้ามเดาค่าที่ไม่มีในโจทย์ เขียน "ข้อมูลไม่เพียงพอ" ถ้าไม่พอ
5. แสดงวิธีคำนวณทีละขั้น ตอบภาษาไทย"""

            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=2048,
                messages=[{"role": "user", "content": prompt}]
            )
            full_answer = message.content[0].text
            answer = clean_for_speech(full_answer[:500])
            docx_file = create_exercise_report(file.filename, full_answer) if want_word else None

        elif want_financial and ext == "csv":
            calc = calculate_from_csv(content)
            prompt = f"""คุณคือ Jarvis AI ผู้เชี่ยวชาญด้านบัญชี
{knowledge}

ข้อมูลคำนวณแล้ว:
- รายได้สุทธิ: {calc['total_income']:,.0f} บาท
- ค่าใช้จ่ายสุทธิ: {calc['total_expense']:,.0f} บาท
- กำไรสุทธิ: {calc['net_profit']:,.0f} บาท

ข้อมูลดิบ:
{content[:4000]}

คำนวณและตรวจทานในใจให้เสร็จก่อน แล้วตอบกลับมาเป็น JSON object เพียงอย่างเดียวเท่านั้น ห้ามมีข้อความ คำอธิบาย หรือขั้นตอนการคิดใดๆ ปนอยู่ก่อนหรือหลัง JSON เด็ดขาด (ฟิลด์ "summary" ต้องเขียนเป็นภาษาไทยเสมอ):
{{"current_assets":0,"current_liabilities":0,"equity":0,"cash_in":0,"cash_out":0,"net_cash":0,"summary":"<วิเคราะห์>"}}"""

            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}]
            )
            raw = message.content[0].text
            try:
                ai_data = extract_json(raw)
                ai_data.update(calc)
                answer = clean_for_speech(ai_data.get("summary", ""))
                docx_file = create_excel_report(file.filename, ai_data, answer)
            except Exception as e:
                print(f"JSON parse error: {e}")
                answer = "ขออภัยครับ Jarvis ประมวลผลข้อมูลไม่สำเร็จ ลองส่งคำสั่งใหม่อีกครั้งครับ"
                docx_file = None

        else:
            prompt = f"""คุณคือ Jarvis AI
คำสั่ง: {command}
ข้อมูล {ext.upper()}:
{content[:6000]}
ตอบภาษาไทย ประโยคธรรมชาติ"""
            message = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}]
            )
            answer = clean_for_speech(message.content[0].text)
            docx_file = None

    except Exception as e:
        print(f"Anthropic API error: {e}")
        return jsonify({"error": f"เรียก AI ไม่สำเร็จ: {e}"}), 500

    write_state(answer, file.filename)

    meta = ai_data if ai_data else {}
    return jsonify({"result": answer, "filename": file.filename, "docx_file": docx_file, "meta": meta})


@app.route("/analyze-image", methods=["POST"])
def analyze_image():
    file = request.files.get("image")
    command = request.form.get("command", "วิเคราะห์รูปภาพนี้")
    if not file:
        return jsonify({"error": "no image"}), 400

    img_data = base64.standard_b64encode(file.read()).decode("utf-8")
    ext = file.filename.lower().split(".")[-1]
    media_type_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "gif": "image/gif", "webp": "image/webp"
    }
    media_type = media_type_map.get(ext, "image/jpeg")

    prompt = f"""คุณคือ Jarvis AI ผู้ช่วยส่วนตัว วิเคราะห์รูปภาพนี้แล้วตอบตามคำสั่ง

คำสั่ง: {command}

ถ้าเป็นใบเสร็จหรือเอกสารการเงิน ให้สรุป:
1. รายการสินค้า/บริการและราคา
2. ยอดรวม
3. วันที่และร้านค้า

ถ้าเป็นรูปทั่วไป ให้อธิบายสิ่งที่เห็น
ตอบภาษาไทย ประโยคธรรมชาติ ไม่ใช้สัญลักษณ์"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": img_data
                        }
                    },
                    {"type": "text", "text": prompt}
                ]
            }]
        )
        answer = clean_for_speech(message.content[0].text)
    except Exception as e:
        print(f"Anthropic API error: {e}")
        return jsonify({"error": f"เรียก AI ไม่สำเร็จ: {e}"}), 500

    write_state(answer, file.filename)

    return jsonify({"result": answer, "filename": file.filename})


@app.route("/download")
def download():
    fname = request.args.get("file")
    if not fname:
        return jsonify({"error": "ไม่ระบุไฟล์"}), 400
    path = os.path.join(OUTPUT_DIR, fname)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    return jsonify({"error": "ไม่พบไฟล์"}), 404


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)